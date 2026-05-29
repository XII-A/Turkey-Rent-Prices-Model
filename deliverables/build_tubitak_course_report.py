from __future__ import annotations

import csv
import json
from collections import Counter, defaultdict
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "House_Rental_Price_Estimation_Tubitak_Style_Course_Report.docx"
CHART_PATH = OUT_DIR / "model_comparison.png"

RAW_DATA = ROOT / "final_data_v4.csv"
MODEL_DATA = ROOT / "finalDataModel.csv"
METRICS_PATH = ROOT / "model_metrics.json"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE = "F4F8FB"
MID_GREY = "666666"
TEXT = "1F1F1F"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "B7C9D6", size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for edge, value in values.items():
        tag = "w:{}".format(edge)
        element = margins.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_keep_together(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_run_font(run, name: str = "Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)


def style_paragraph(paragraph, size=9, color=TEXT, bold=False, italic=False, align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(5)
    paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
        run.italic = italic


def add_paragraph(doc: Document, text: str = "", size=9, color=TEXT, bold=False, italic=False, align=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_paragraph(paragraph, size=size, color=color, bold=bold, italic=italic, align=align)
    return paragraph


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.font.size = Pt(12 if level == 1 else 10)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(5)
    if level == 1:
        paragraph.paragraph_format.keep_with_next = True
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), BLUE)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
    return paragraph


def add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    set_cell_border(cell, "AFC5D6")
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    p.add_run(title + " ").bold = True
    p.add_run(body)
    style_paragraph(p, size=9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.18)
        run = p.add_run("- " + item)
        style_paragraph(p, size=9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        set_cell_border(hdr[idx])
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        style_paragraph(hdr[idx].paragraphs[0], size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(cells[idx].paragraphs[0], size=8.2, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    set_table_keep_together(table)
    add_paragraph(doc, "", size=3)
    return table


def read_metrics() -> dict:
    with METRICS_PATH.open(encoding="utf-8") as fh:
        return json.load(fh)


def summarize_csvs() -> dict:
    raw_rows = 0
    raw_headers = []
    city_counts = Counter()
    room_counts = Counter()
    heating_counts = Counter()
    prices_by_city: dict[str, list[int]] = defaultdict(list)
    prices: list[int] = []
    with RAW_DATA.open(encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        raw_headers = reader.fieldnames or []
        for row in reader:
            raw_rows += 1
            city_counts[row.get("city", "")] += 1
            room_counts[row.get("Oda Sayısı", "")] += 1
            heating_counts[row.get("Isıtma Tipi", "")] += 1
            try:
                price = int((row.get("fiyat", "") or "").replace("TL", "").replace(",", "").replace(".", ""))
                prices.append(price)
                prices_by_city[row.get("city", "")].append(price)
            except ValueError:
                pass

    model_rows = 0
    model_headers = []
    numeric = defaultdict(list)
    with MODEL_DATA.open(encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        model_headers = reader.fieldnames or []
        for row in reader:
            model_rows += 1
            for col in ["fiyat", "Net Metrekare", "Brüt Metrekare", "Binanın Yaşı", "Banyo Sayısı"]:
                try:
                    if row.get(col, "") != "":
                        numeric[col].append(float(row[col]))
                except ValueError:
                    pass

    def median(values: list[float]) -> float:
        values = sorted(values)
        if not values:
            return 0.0
        n = len(values)
        mid = n // 2
        return values[mid] if n % 2 else (values[mid - 1] + values[mid]) / 2

    city_summary = []
    for city, values in prices_by_city.items():
        if values:
            city_summary.append((city, len(values), round(median(values)), round(sum(values) / len(values))))
    city_summary.sort(key=lambda item: item[1], reverse=True)

    return {
        "raw_rows": raw_rows,
        "raw_columns": len(raw_headers),
        "model_rows": model_rows,
        "model_columns": len(model_headers),
        "cities_total": len([c for c in city_counts if c]),
        "top_cities": city_counts.most_common(8),
        "top_rooms": room_counts.most_common(6),
        "top_heating": heating_counts.most_common(5),
        "price_median": round(median(prices)),
        "price_mean": round(sum(prices) / len(prices)) if prices else 0,
        "model_medians": {col: round(median(values), 2) for col, values in numeric.items()},
        "city_price_summary": city_summary[:8],
    }


def make_chart(metrics: dict) -> None:
    labels = [item["model"].replace("Artificial Neural Network (Keras)", "Keras ANN") for item in metrics["results"]]
    rmse = [item["rmse"] for item in metrics["results"]]
    r2 = [item["r2"] for item in metrics["results"]]
    x = range(len(labels))
    fig, ax1 = plt.subplots(figsize=(7.2, 3.6), dpi=180)
    ax1.bar([i - 0.18 for i in x], rmse, width=0.35, label="RMSE", color="#3D7EA6")
    ax1.set_ylabel("RMSE on log1p rent")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(labels, rotation=0, ha="center")
    ax1.set_ylim(0, max(rmse) * 1.25)
    ax1.grid(axis="y", color="#D6DEE6", linewidth=0.6)
    ax2 = ax1.twinx()
    ax2.bar([i + 0.18 for i in x], r2, width=0.35, label="R2", color="#78A083")
    ax2.set_ylabel("R2")
    ax2.set_ylim(0, max(r2) * 1.35)
    lines, labels_a = ax1.get_legend_handles_labels()
    lines2, labels_b = ax2.get_legend_handles_labels()
    ax1.legend(lines + lines2, labels_a + labels_b, loc="upper right", frameon=False)
    ax1.set_title("Reproducible Model Comparison")
    fig.tight_layout()
    fig.savefig(CHART_PATH, bbox_inches="tight")
    plt.close(fig)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "1002BBF-01 style course adaptation | FBE500i | Abdulrazak Armanazi"
    style_paragraph(footer_para, size=8, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)


def add_cover(doc: Document) -> None:
    add_paragraph(doc, "1002 - HIZLI DESTEK PROGRAMI", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "1002-A HIZLI DESTEK MODÜLÜ", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "PROJE BAŞVURU FORMU", size=16, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Course Report Adaptation", size=10, italic=True, color=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "", size=8)

    rows = [
        ["Project Title", "House Rental Price Estimation in Türkiye"],
        ["Researcher", "Abdulrazak Armanazi (Student No: Y255050086)"],
        ["Institution / University", "Sakarya University"],
        ["Faculty / Institute", "FEN BİLİMLERİ ENSTİTÜSÜ"],
        ["Department / Program", "YAZILIM MÜHENDİSLİĞİ (YL) (TEZLİ) (İNGİLİZCE)"],
        ["Course", "FBE500i Scientific Research, Project Writing Techniques and Seminar"],
        ["Supervisor / Advisor", "AHMET ÖZMEN"],
        ["Project Period", "27/03/2026 - 26/05/2026"],
        ["Submission Date", "26/05/2026"],
    ]
    add_table(doc, ["Field", "Information"], rows, widths=[4.4, 11.6])
    add_note_box(
        doc,
        "Document note:",
        "This file adapts the TÜBİTAK 1002-A project application structure for a completed FBE500i course report. "
        "The report uses the provided slide deck and the reproducible codebase as the primary evidence base.",
    )
    doc.add_page_break()


def add_scientific_quality(doc: Document, stats: dict) -> None:
    add_heading(doc, "1. BİLİMSEL NİTELİK", 1)
    add_heading(doc, "Konunun Önemi ve Projenin Bilimsel Niteliği", 2)
    paragraphs = [
        "Residential rent is one of the most visible economic signals for students, families, employees, property owners, and policy observers. "
        "In Türkiye, rental listings differ sharply by city, district, housing size, room structure, heating system, building attributes, balcony availability, and local market conditions. "
        "Because these factors interact with each other, a rule-of-thumb estimate is often too weak to explain why similar-looking homes can have different advertised rents.",
        "The project addresses this problem as a supervised regression task: estimate continuous monthly rent from observable listing attributes. "
        "The scientific value of the work is not only the final prediction score, but the complete pipeline that turns messy public listing data into a reproducible modeling dataset. "
        "The codebase documents data acquisition, preprocessing, feature engineering, benchmark model training, and metric reporting in a way that can be repeated and improved.",
        f"The available raw source file contains {stats['raw_rows']:,} rental listings and {stats['raw_columns']} columns. "
        f"The modeling file contains {stats['model_rows']:,} rows and {stats['model_columns']} columns before the model script removes missing values and price outliers. "
        f"The current reproducible model run uses {stats['cities_total']} cities, a log-transformed rent target, and {stats['model_medians'].get('Net Metrekare', 0):.0f} m2 as the median net area in the model-ready data.",
        "The research question is: To what extent can structured rental-listing attributes estimate residential rental prices in Türkiye using reproducible machine learning models? "
        "A practical hypothesis follows from the real-estate valuation literature: location, property size, room count, and building amenities should explain a meaningful share of rent variation, while non-linear models may capture interactions better than a purely linear baseline.",
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
    add_bullets(
        doc,
        [
            "Target variable: monthly rent, modeled as log1p(fiyat) to reduce the influence of extreme price values.",
            "Independent variables: net/gross area, floor and building information, heating type, site/balcony status, room counts, bathroom and WC counts, latitude, longitude, and related encoded listing attributes.",
            "Scientific contribution in the course context: a complete, reproducible case study of data collection, preprocessing, benchmark modeling, and interpretation for the Turkish rental market.",
        ],
    )

    add_heading(doc, "Literature Context and Project Gap", 2)
    for text in [
        "The project is aligned with the hedonic-pricing idea that real-estate value can be approximated from property characteristics and location. "
        "Recent machine learning studies extend that tradition by comparing linear models with decision trees, ensembles, neural networks, and spatial features. "
        "The cited studies do not make the course project novel by themselves; instead, they show that the project follows a recognizable research pattern: structured real-estate attributes can be treated as explanatory variables, while model comparison reveals how much predictive signal is available.",
        "The local gap addressed in this report is the construction of a reproducible student-level pipeline for Turkish rental listings. "
        "Many public examples focus on sale prices, international datasets, or already-clean benchmark data. "
        "This project begins earlier in the pipeline, at public listing extraction and messy attribute conversion, so the final report can discuss the full path from raw listings to benchmark model results.",
        "The scientific boundary is also clear. The model estimates advertised rent from listings; it does not estimate final contract rent, affordability, or causal effects of policy variables. "
        "This boundary is important because listing prices can reflect seller expectations, short-term market pressure, data-entry errors, and temporary local shocks. "
        "For a course project, the appropriate claim is therefore methodological and educational: the workflow demonstrates how a Turkish rental-price dataset can be transformed into a supervised learning problem and evaluated transparently.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "Amaç ve Hedefler", 2)
    for text in [
        "The general aim of the project is to build and evaluate a reproducible machine learning workflow for estimating residential rental prices in Türkiye from public listing attributes.",
        "The first objective is to collect and structure rental listing information from a real estate portal. "
        "The second objective is to clean the data, convert text-based variables into numerical features, and engineer additional variables such as total room count and geolocation. "
        "The third objective is to compare baseline and non-linear models under the same train/test split. "
        "The final objective is to report the model behavior clearly enough that another student or researcher can reproduce the experiment and identify improvement paths.",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        ["Objective", "Measurable Output"],
        [
            ["Data collection", "Structured source dataset from public rental listings."],
            ["Preprocessing", "Numeric model-ready dataset with missing-value handling, encoding, scaling, and outlier filtering."],
            ["Modeling", "Linear Regression, Decision Tree, and Keras ANN trained with the same reproducible split."],
            ["Evaluation", "MSE, MAE, RMSE, R2, and explained variance recorded in model_metrics.json."],
            ["Reporting", "Course report explaining motivation, method, results, limitations, and future work."],
        ],
        widths=[4.4, 11.6],
    )
    add_paragraph(
        doc,
        "The expected learning outcome is a stronger understanding of how project-writing structure, data science workflow, and reproducible evaluation connect. "
        "The report also turns the slide deck into a more complete written record by documenting the evidence behind each slide: where the data came from, what transformations were applied, which models were compared, and what the current benchmark metrics imply.",
    )


def add_method(doc: Document, metrics: dict, stats: dict) -> None:
    add_heading(doc, "2. YÖNTEM", 1)
    for text in [
        "The method follows the workflow described in the slides and implemented in the repository: web scraping, preprocessing, feature engineering, model training, evaluation, and reporting. "
        "The design is observational and data-driven. It does not perform an experiment on human participants; instead, it analyzes public listing attributes as predictors of advertised rent.",
        "The original scraper visits Emlakjet rental-listing pages by city, extracts listing links, opens each listing detail page, and records price, city, district, neighborhood, and available property attributes. "
        "The code comments state that the scraper should not be re-run unless the dataset needs to be refreshed, because the current CSV files already contain the collected data.",
        "Preprocessing converts rent strings to integers, converts square-meter fields to numeric form, splits room-count text into bedroom and living-room variables, encodes binary fields such as site and balcony status, frequency-encodes selected categorical fields, and adds latitude/longitude fields from geocoding. "
        "The notebook also investigates outliers and uses z-score based filtering on price before modeling.",
        "The final model script uses finalDataModel.csv, drops missing rows, applies log1p to the rent target, filters extreme target values using a z-score threshold, combines bedrooms and living rooms into a total rooms feature, and scales predictors with MinMaxScaler. "
        "The train/test split is 70/30 with random_state=101.",
    ]:
        add_paragraph(doc, text)

    add_table(
        doc,
        ["Dataset / Artifact", "Role in the Method"],
        [
            ["webscraper.py", "Collects public rental listing details from Emlakjet pages."],
            ["preprocessing.ipynb", "Cleans, converts, encodes, geocodes, and prepares feature columns."],
            ["final_data_v4.csv", f"Raw structured export with {stats['raw_rows']:,} rows and {stats['raw_columns']} columns."],
            ["finalDataModel.csv", f"Model-ready numeric dataset with {stats['model_rows']:,} rows and {stats['model_columns']} columns."],
            ["model.py", "Trains benchmark regressors and writes reproducible metrics."],
            ["model_metrics.json", "Stores the final performance table used in this report."],
        ],
        widths=[4.2, 11.8],
    )

    add_heading(doc, "Data Characteristics", 2)
    add_table(
        doc,
        ["Indicator", "Observed Value"],
        [
            ["Top listed cities", ", ".join([f"{city} ({count})" for city, count in stats["top_cities"][:5]])],
            ["Most common room counts", ", ".join([f"{room} ({count})" for room, count in stats["top_rooms"][:5] if room])],
            ["Most common heating types", ", ".join([f"{heat} ({count})" for heat, count in stats["top_heating"][:4] if heat])],
            ["Median raw listed rent", f"{stats['price_median']:,} TL"],
            ["Median net / gross area", f"{stats['model_medians'].get('Net Metrekare', 0):.0f} m2 / {stats['model_medians'].get('Brüt Metrekare', 0):.0f} m2"],
        ],
        widths=[4.5, 11.5],
    )

    add_heading(doc, "Preprocessing and Feature Engineering Details", 2)
    for text in [
        "The preprocessing notebook shows that most of the project effort is concentrated before model fitting. "
        "Raw listing exports include fields that are useful for describing a property, fields that are mostly empty, and fields that are administrative rather than predictive. "
        "The notebook removes columns such as listing number, listing dates, category labels, and very sparse balcony-type variants when they do not add stable predictive value for the benchmark.",
        "Price values are cleaned by removing the TL marker and thousands separators, then converting to integer form. "
        "Net and gross square-meter fields are converted from text to floating-point values. "
        "Room-count values are normalized by replacing special labels such as 'Stüdyo' and splitting the room-count string into bedroom and living-room variables. "
        "The model script later combines these into a total room count to reduce redundancy.",
        "Categorical variables are handled with pragmatic encodings suitable for a first reproducible benchmark. "
        "Binary indicators such as 'Site İçerisinde' and 'Balkon Durumu' are mapped to numeric values. "
        "Heating type is frequency encoded after rare categories are grouped, which preserves information about common heating systems without exploding the feature space into many sparse dummy variables. "
        "Geolocation is represented through latitude and longitude fields, allowing the model to use coarse spatial signal without needing a complex geographic hierarchy.",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        ["Preprocessing Step", "Reason", "Effect on Modeling"],
        [
            ["Remove low-value / sparse columns", "Avoid noise from mostly empty or administrative fields.", "Keeps the feature set compact and easier to reproduce."],
            ["Clean rent and area fields", "Raw strings contain currency and unit text.", "Makes numerical scaling and regression possible."],
            ["Room-count conversion", "Listing text contains mixed room formats.", "Creates bedroom, living-room, and total-room predictors."],
            ["Binary and frequency encoding", "Models require numeric inputs.", "Preserves amenity and heating signals in model-ready form."],
            ["Outlier filtering", "Extreme values can dominate squared-error objectives.", "Stabilizes comparison across benchmark models."],
            ["MinMax scaling", "The ANN and linear model benefit from comparable feature ranges.", "Improves training stability and fairer model comparison."],
        ],
        widths=[4.2, 6.0, 5.8],
    )

    add_heading(doc, "Modeling and Evaluation Strategy", 2)
    add_paragraph(
        doc,
        "Three benchmark models are used so that model behavior can be compared from simple to more flexible: Linear Regression, Decision Tree Regressor, and an Artificial Neural Network implemented with Keras. "
        "Linear Regression provides an interpretable baseline. The Decision Tree captures non-linear splits in the feature space. The ANN tests whether a deeper function approximator improves predictive accuracy after scaling.",
    )
    add_paragraph(
        doc,
        "The ANN architecture uses Dense layers with ReLU activations, L2 regularization, Batch Normalization, Dropout, Adam optimizer, and EarlyStopping on validation loss. "
        "All models predict log1p rent, and the evaluation metrics are computed on the held-out test set. Because the target is log-transformed, errors should be interpreted as model comparison signals rather than direct Turkish lira errors.",
    )
    rows = []
    for item in metrics["results"]:
        rows.append(
            [
                item["model"],
                f"{item['mse']:.4f}",
                f"{item['mae']:.4f}",
                f"{item['rmse']:.4f}",
                f"{item['r2']:.4f}",
            ]
        )
    add_table(doc, ["Model", "MSE", "MAE", "RMSE", "R2"], rows, widths=[5.6, 2.4, 2.4, 2.4, 2.4])
    doc.add_picture(str(CHART_PATH), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        doc,
        "Figure 1. Reproducible benchmark comparison from model_metrics.json. The Decision Tree model has the lowest RMSE and highest R2 among the current benchmark models.",
        size=8,
        italic=True,
        color=MID_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_heading(doc, "Result Interpretation", 2)
    best = min(metrics["results"], key=lambda item: item["rmse"])
    for text in [
        f"The strongest reproducible benchmark is {best['model']}, with RMSE {best['rmse']:.4f} and R2 {best['r2']:.4f}. "
        "This means the tree-based model explains more held-out variation than the linear baseline and the current neural-network configuration. "
        "The result is plausible for this dataset because rental prices often change through threshold-like effects: particular cities, property sizes, room structures, or building characteristics may create non-linear splits that a simple linear equation cannot capture as easily.",
        "The ANN does improve over the linear model on R2, but it does not outperform the Decision Tree in the current run. "
        "This does not mean neural networks are unsuitable for rent estimation in general. "
        "It means that the current architecture, feature set, dataset size, and preprocessing choices do not yet give the ANN a clear advantage over the simpler tree model. "
        "For a completed course report, that is a useful finding because it discourages assuming that a more complex model is automatically better.",
        "The R2 values also show that the current feature set leaves unexplained variation. "
        "Advertised rent can depend on micro-location, building condition, furnishing quality, view, distance to transport, seasonal timing, negotiation behavior, and macroeconomic conditions. "
        "Many of these signals are missing or only indirectly represented in the current listing fields. "
        "Therefore, the benchmark should be interpreted as a reproducible starting point rather than a finished commercial valuation tool.",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        ["Model", "Interpretation in This Project"],
        [
            ["Linear Regression", "Useful baseline; limited because rent-price relationships are not purely additive or linear."],
            ["Decision Tree", "Best current benchmark; captures non-linear thresholds and interactions in listing attributes."],
            ["Keras ANN", "More flexible architecture, but current feature set and configuration do not outperform the tree model."],
        ],
        widths=[4.2, 11.8],
    )
    doc.add_page_break()


def add_management(doc: Document) -> None:
    add_heading(doc, "3. PROJE YÖNETİMİ", 1)
    add_heading(doc, "3.1. İş-Zaman Çizelgesi", 2)
    add_paragraph(
        doc,
        "The project was completed as a course project between 27/03/2026 and 26/05/2026. "
        "The schedule below is intentionally simple because the work was carried out by one researcher under course supervision rather than by a large funded project team.",
    )
    add_table(
        doc,
        ["IP No", "Work Package", "Responsible Person", "Period", "Main Output"],
        [
            ["1", "Project framing and dataset review", "Abdulrazak Armanazi", "27/03-05/04/2026", "Problem statement and data requirements."],
            ["2", "Data collection and source structuring", "Abdulrazak Armanazi", "06/04-17/04/2026", "Structured rental listing CSV."],
            ["3", "Preprocessing and feature engineering", "Abdulrazak Armanazi", "18/04-30/04/2026", "Model-ready numeric dataset."],
            ["4", "Model training and metric comparison", "Abdulrazak Armanazi", "01/05-15/05/2026", "Linear, Decision Tree, and ANN benchmark results."],
            ["5", "Presentation and final report preparation", "Abdulrazak Armanazi", "16/05-26/05/2026", "Slides and course report."],
        ],
        widths=[1.3, 4.2, 3.7, 3.0, 4.0],
    )
    add_heading(doc, "3.2. Başarı Ölçütleri", 2)
    add_table(
        doc,
        ["IP No", "Success Criterion", "Contribution (%)"],
        [
            ["1", "A clear research problem and scope are defined for Turkish residential rent estimation.", "15"],
            ["2", "A structured dataset is available with price, location, physical, and amenity variables.", "20"],
            ["3", "Raw attributes are converted into a numeric modeling dataset with engineered features.", "25"],
            ["4", "At least three models are trained and evaluated using reproducible metrics.", "25"],
            ["5", "Findings, limitations, and future work are communicated in slides and report format.", "15"],
        ],
        widths=[1.5, 12.0, 2.5],
    )
    add_heading(doc, "3.3. Risk Yönetimi", 2)
    add_paragraph(
        doc,
        "No major unresolved risks remain because the work has already been completed for the course context. "
        "For completeness, the table lists routine technical limitations that would matter if the project were repeated or expanded.",
    )
    add_table(
        doc,
        ["IP No", "Routine Limitation", "Mitigation / B Plan"],
        [
            ["2", "Public listing pages can change structure or restrict scraping.", "Use the saved CSV for reproducibility; update scraper selectors only if a new data refresh is required."],
            ["3", "Some listing attributes are missing or inconsistent.", "Drop unusable columns, encode available fields consistently, and report missing-value handling transparently."],
            ["4", "Flexible models may overfit listing noise.", "Use train/test split, regularization, early stopping, and compare against simpler baselines."],
        ],
        widths=[1.4, 6.2, 8.4],
    )
    add_paragraph(
        doc,
        "The risk plan is deliberately brief because the assignment has already reached completion. "
        "However, the listed limitations are still useful for academic transparency: they explain why the saved datasets and reproducible script should be treated as the authoritative evidence for this submission, and why future extensions should be planned as new work rather than silently changing the submitted result.",
    )


def add_outputs_budget_references(doc: Document, metrics: dict) -> None:
    add_heading(doc, "4. ÇIKTI, ETKİ VE KAZANIMLAR", 1)
    add_paragraph(
        doc,
        "The project produced an end-to-end workflow for estimating residential rents in Türkiye. "
        "The immediate output is a reproducible educational machine learning project; the broader benefit is a clearer understanding of how property attributes and location features can be organized for rent estimation.",
    )
    add_table(
        doc,
        ["Category", "Expected / Achieved Output", "Beneficiaries", "Time"],
        [
            ["Scientific / educational output", "A completed course report, slide deck, preprocessing notebook, model script, and reproducible metrics file.", "Student, supervisor, and course reviewers.", "0-6 months"],
            ["Technical output", "A benchmark rent-estimation pipeline using Linear Regression, Decision Tree, and Keras ANN.", "Future students extending the project.", "0-6 months"],
            ["Social / practical effect", "A transparent example of data-driven rental-price estimation that can support fairer price discussions.", "Renters, landlords, and researchers in an educational context.", "Project sonrası"],
            ["Future research", "Possible extension with inflation data, richer location variables, larger data refreshes, and stronger ensemble models.", "Researchers and graduate students.", "Project sonrası"],
        ],
        widths=[3.3, 6.5, 4.3, 2.0],
    )

    add_paragraph(
        doc,
        "The strongest current benchmark is the Decision Tree model, with MSE "
        f"{metrics['results'][1]['mse']:.4f}, RMSE {metrics['results'][1]['rmse']:.4f}, and R2 {metrics['results'][1]['r2']:.4f}. "
        "The result is useful as a baseline, but it also shows that rent estimation remains difficult: many price drivers are local, temporal, and sometimes unavailable in listing data. "
        "Therefore, the report treats the model as an educational and reproducible benchmark rather than a production valuation system.",
    )

    add_heading(doc, "Completed Course Outcomes", 2)
    for text in [
        "In relation to the FBE500i course, the project demonstrates the main stages of scientific project writing: identifying a problem, motivating the research question, defining data and method, producing measurable outputs, and reflecting on limitations. "
        "The slide deck provides the presentation form of the work, while this report provides the fuller written justification and technical explanation.",
        "The project also shows the value of reproducibility. The README explicitly documents the workflow and the benchmark table, while model.py writes a machine-readable metrics file. "
        "This makes the submitted results easier to audit than a report that only states final scores. "
        "If another reviewer runs the same script in a compatible environment, the expected output structure and evaluation metrics are already specified.",
        "From a learning perspective, the most important gain is not a single R2 score. The main gain is understanding how each decision in the pipeline affects the credibility of the final result: scraped source data determines coverage, preprocessing determines the usable feature set, model choice determines the hypothesis tested, and evaluation determines whether the result is persuasive.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "Future Work", 2)
    add_bullets(
        doc,
        [
            "Refresh the data periodically and add a clear collection date so the model can reflect market changes over time.",
            "Add richer location variables such as distance to city center, transit access, university districts, coastal areas, or neighborhood-level median rent.",
            "Compare stronger ensemble models such as Random Forest, Gradient Boosting, XGBoost, or LightGBM under the same split and reporting format.",
            "Evaluate predictions after converting log-scale errors back into approximate Turkish lira differences for easier practical interpretation.",
            "Develop a small user interface only after the modeling pipeline has stronger validation and clearer data-refresh rules.",
        ],
    )

    add_heading(doc, "BELİRTMEK İSTEDİĞİNİZ DİĞER KONULAR", 1)
    add_paragraph(
        doc,
        "The original TÜBİTAK form is intended for proposed research funding. This document intentionally adapts that structure to a completed FBE500i course report. "
        "The project therefore uses the form's scientific-quality, method, management, output, and budget sections as an academic reporting framework.",
    )

    add_heading(doc, "BAŞVURU FORMU EKLERİ", 1)
    add_heading(doc, "EK-1: KAYNAKLAR", 2)
    references = [
        "TÜBİTAK. (2026). 1002-A Hızlı Destek Modülü. Türkiye Bilimsel ve Teknolojik Araştırma Kurumu. https://tubitak.gov.tr/tr/icerik-1002-a-hizli-destek-modulu",
        "Çılgın, C., & Gökçen, H. (2023). Machine learning methods for prediction real estate sales prices in Turkey. Revista de la Construcción, 22(1), 163-177.",
        "Seya, H., & Shiroi, D. (2021). Spatial prediction of apartment rent using regression-based and machine learning-based approaches with a large dataset. arXiv:2107.12539.",
        "Cao, Y., Ma, J., & Chen, Y. (2024). Research on renting price prediction based on machine learning. EUDL. https://doi.org/10.4108/eai.8-12-2023.2344718",
        "Armanazi, A. (2026). House Rental Price Estimation in Türkiye project codebase: webscraper.py, preprocessing.ipynb, model.py, final_data_v4.csv, finalDataModel.csv, and model_metrics.json.",
    ]
    for ref in references:
        p = add_paragraph(doc, ref, size=8.5)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)

    add_heading(doc, "EK-2: BÜTÇE VE GEREKÇESİ", 2)
    add_paragraph(
        doc,
        "Because this is a completed course report adaptation, the budget is presented as a simple record of basic resources used rather than as a formal funding request.",
    )
    add_table(
        doc,
        ["Item", "Purpose", "Estimated Cost / Status"],
        [
            ["Laptop / personal computer", "Data cleaning, notebook execution, model training, presentation, and report writing.", "Existing student resource; no purchase requested."],
            ["Internet subscription", "Access to public listing pages, documentation, references, and course materials.", "Existing monthly service; no separate project purchase requested."],
        ],
        widths=[4.2, 8.0, 3.8],
    )


def build() -> None:
    metrics = read_metrics()
    stats = summarize_csvs()
    make_chart(metrics)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_scientific_quality(doc, stats)
    add_method(doc, metrics, stats)
    add_management(doc)
    add_outputs_budget_references(doc, metrics)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
