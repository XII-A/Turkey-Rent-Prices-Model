from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "EK-1_Kaynaklar.docx"

BLUE = "1F4E79"
GREY = "666666"
TEXT = "1F1F1F"


def set_run_font(run, name: str = "Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)


def style_paragraph(paragraph, size=9, color=TEXT, bold=False, italic=False, align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
        run.italic = italic


def add_paragraph(doc: Document, text: str = "", size=9, color=TEXT, bold=False, italic=False, align=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_paragraph(paragraph, size=size, color=color, bold=bold, italic=italic, align=align)
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.paragraph_format.space_after = Pt(8)

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.text = "1002BF-01 | EK-1: KAYNAKLAR | Abdulrazak Armanazi"
    style_paragraph(footer, size=8, color=GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)


def add_reference(doc: Document, text: str) -> None:
    paragraph = add_paragraph(doc, text, size=9)
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.75)
    paragraph.paragraph_format.space_after = Pt(7)


def build() -> None:
    doc = Document()
    configure_document(doc)

    add_heading(doc, "EK-1: KAYNAKLAR")
    add_paragraph(
        doc,
        "Bu bölümde, proje önerisinde yararlanılan kaynakların listesi TÜBİTAK Bibliyografik Verilerin "
        "Düzenlenmesi sayfasındaki açıklamalara uygun olarak verilmeli ve bu kaynaklara metnin içerisinde "
        "atıf yapılmalıdır.",
        size=9,
    )
    add_paragraph(
        doc,
        "Kaynaklar yazar soyadlarına / kurum adına göre alfabetik sırada düzenlenmiş; DOI numarası bulunan "
        "kaynaklarda DOI bilgisi kaynak sonuna eklenmiş ve internet kaynakları için son erişim tarihi ayrıca "
        "belirtilmiştir.",
        size=9,
        italic=True,
        color=GREY,
    )
    add_paragraph(doc, "", size=4)

    references = [
        (
            "Armanazi, A. 2026. “House Rental Price Estimation in Türkiye Project Codebase and Datasets”, "
            "Sakarya University FBE500i Course Project Files, webscraper.py, preprocessing.ipynb, model.py, "
            "final_data_v4.csv, finalDataModel.csv, model_metrics.json."
        ),
        (
            "Cao, S., Liao, W., Huang, J. 2024. “Research on Renting Price Prediction Based on Machine "
            "Learning”, Proceedings of the 5th Management Science Informatization and Economic Innovation "
            "Development Conference, MSIEID 2023. DOI: 10.4108/eai.8-12-2023.2344718."
        ),
        (
            "Çılgın, C., Gökçen, H. 2023. “Machine Learning Methods for Prediction Real Estate Sales Prices "
            "in Turkey”, Revista de la Construcción. Journal of Construction, 22(1), 163-177. "
            "DOI: 10.7764/RDLC.22.1.163."
        ),
        (
            "Lu, S., Li, Z., Qin, Z., Yang, X., Goh, R. S. M. 2017. “A Hybrid Regression Technique for "
            "House Prices Prediction”, 2017 IEEE International Conference on Industrial Engineering and "
            "Engineering Management (IEEM), Singapore, 319-323. DOI: 10.1109/IEEM.2017.8289904."
        ),
        (
            "Phan, T. D. 2018. “Housing Price Prediction Using Machine Learning Algorithms: The Case of "
            "Melbourne City, Australia”, 2018 International Conference on Machine Learning and Data "
            "Engineering (iCMLDE), Sydney, NSW, Australia, 35-42. DOI: 10.1109/iCMLDE.2018.00017."
        ),
        (
            "Selim, H. 2009. “Determinants of House Prices in Turkey: Hedonic Regression versus Artificial "
            "Neural Network”, Expert Systems with Applications, 36(2), 2843-2852. "
            "DOI: 10.1016/J.ESWA.2008.01.044."
        ),
        (
            "Türkiye Bilimsel ve Teknolojik Araştırma Kurumu. “1002-A Hızlı Destek Modülü”. "
            "https://tubitak.gov.tr/tr/destekler/akademik/ulusal-destek-programlari/1002-hizli-destek-modulu, "
            "Son erişim tarihi: 29 Mayıs 2026."
        ),
        (
            "Yoshida, T., Murakami, D., Seya, H. 2022. “Spatial Prediction of Apartment Rent Using "
            "Regression-Based and Machine Learning-Based Approaches with a Large Dataset”, "
            "The Journal of Real Estate Finance and Economics, 69(1), 1-28. "
            "DOI: 10.1007/s11146-022-09929-6."
        ),
    ]

    for reference in references:
        add_reference(doc, reference)

    add_paragraph(
        doc,
        "Not: Kaynakça biçimi TÜBİTAK’ın Bibliyografik Verilerin Düzenlenmesi duyurusundaki periyodik yayın "
        "ve internet kaynağı örnekleri esas alınarak hazırlanmıştır.",
        size=8,
        italic=True,
        color=GREY,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
